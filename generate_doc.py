# -*- coding: utf-8 -*-
from docx import Document  
from docx.shared import  Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT , WD_ALIGN_PARAGRAPH , WD_TAB_ALIGNMENT
from PySide6.QtWidgets import  QFileDialog, QMessageBox
from load_constants import  categorys , months , city , ecole
from docx.oxml import parse_xml
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime, timedelta , date

def configure_document(document):
    # تحديد الاتجاه RTL للمستند
    pPr = document.styles['Normal']._element.get_or_add_pPr()
    bidi_element = parse_xml('<w:bidi xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
    pPr.append(bidi_element)

    # تحديد XML للنمط الأول (عنوان) مع تحديد اللغة العربية والاتجاه RTL وحجم الخط دون خاصية التسطير
    title_style_xml = """
    <w:style xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
             xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
             xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
             xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml"
             w:type="paragraph"
             w:styleId="Title">
      <w:name w:val="Title"/>
      <w:basedOn w:val="Normal"/>
      <w:next w:val="Normal"/>
      <w:link w:val="TitleChar"/>
      <w:uiPriority w:val="10"/>
      <w:qFormat/>
      <w:rsid w:val="00FC693F"/>
      <w:pPr>
        <w:spacing w:after="300" w:line="240" w:lineRule="auto"/>
        <w:contextualSpacing/>
        <w:bidi/>
      </w:pPr>
      <w:rPr>
        <w:rFonts w:asciiTheme="majorHAnsi" w:eastAsiaTheme="majorEastAsia" w:hAnsiTheme="majorHAnsi" w:cstheme="majorBidi"/>
        <w:color w:val="17365D" w:themeColor="text2" w:themeShade="BF"/>
        <w:spacing w:val="5"/>
        <w:kern w:val="28"/>
        <w:sz w:val="36"/>
        <w:szCs w:val="36"/>
        <w:lang w:val="ar-SA"/>
        <w:bidi/>
      </w:rPr>
    </w:style>
    """

    # تحديد XML للنمط الثاني (فقرة) مع تحديد اللغة العربية والاتجاه RTL وحجم الخط دون خاصية التسطير
    paragraph_style_xml = """
    <w:style xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
             xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
             xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
             xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml"
             w:type="paragraph"
             w:styleId="Paragraph">
      <w:name w:val="Paragraph"/>
      <w:basedOn w:val="Normal"/>
      <w:next w:val="Normal"/>
      <w:uiPriority w:val="99"/>
      <w:unhideWhenUsed/>
      <w:rsid w:val="00FC693F"/>
      <w:pPr>
        <w:spacing w:after="200" w:line="240" w:lineRule="auto"/>
        <w:contextualSpacing/>
        <w:bidi/>
      </w:pPr>
      <w:rPr>
        <w:rFonts w:asciiTheme="majorHAnsi" w:eastAsiaTheme="majorEastAsia" w:hAnsiTheme="majorHAnsi" w:cstheme="majorBidi"/>
        <w:color w:val="17365D"/>
        <w:sz w:val="28"/>
        <w:szCs w:val="28"/>
        <w:lang w:val="ar-SA"/>
        <w:bidi/>
      </w:rPr>
    </w:style>
    """
    table_style_xml = """
    <w:style xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
                xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
                xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml"
                w:type="table"
                w:styleId="TableArabic">
        <w:name w:val="TableArabic"/>
        <w:basedOn w:val="TableNormal"/>
        <w:uiPriority w:val="99"/>
        <w:unhideWhenUsed/>
        <w:rsid w:val="00FC693F"/>
        <w:pPr>
            <w:bidi/>
        </w:pPr>
        <w:rPr>
            <w:rFonts w:asciiTheme="majorHAnsi" w:eastAsiaTheme="majorEastAsia" w:hAnsiTheme="majorHAnsi" w:cstheme="majorBidi"/>
            <w:color w:val="17365D"/>
            <w:sz w:val="24"/>
            <w:lang w:val="ar-SA"/>
            <w:bidi/>
        </w:rPr>
        <w:tblPr>
            <w:bidiVisual/>
            <w:tblBorders>
            <w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>
            <w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>
            <w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>
            <w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>
            <w:insideH w:val="single" w:sz="4" w:space="0" w:color="auto"/>
            <w:insideV w:val="single" w:sz="4" w:space="0" w:color="auto"/>
            </w:tblBorders>
        </w:tblPr>
        <w:tcPr>
            <w:tcBorders>
            <w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>
            <w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>
            <w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>
            <w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>
            <w:insideH w:val="single" w:sz="4" w:space="0" w:color="auto"/>
            <w:insideV w:val="single" w:sz="4" w:space="0" w:color="auto"/>
            </w:tcBorders>
        </w:tcPr>
    </w:style>
    """
    # إضافة النمط الأول (عنوان) إلى المستند
    title_style_element = parse_xml(title_style_xml)
    document.styles._element.append(title_style_element)

    # إضافة النمط الثاني (فقرة) إلى المستند
    paragraph_style_element = parse_xml(paragraph_style_xml)
    document.styles._element.append(paragraph_style_element)
    
    # إضافة النمط الأخير (جدول) إلى المستند
    table_style_element = parse_xml(table_style_xml)
    document.styles._element.append(table_style_element)


# # **************************************************************************************************************************
# # **************************************************************************************************************************
# #                             # generator.py
# # **************************************************************************************************************************
# # **************************************************************************************************************************
Logo_header: str = 'الجمهورية الجزائرية الديمقراطية الشعبية\nوزارة التربية الوطنية'

    
def generate_absence_doc(
    output_doc_name: str,
    month: int,
    year: int,
    employees_absences: dict[str, list[dict[str, str]]]
) -> None:
    """
    Creates an absence report with a table for each category in the given document.
    
    Args:
        document (Document): The Word document to add the tables to.
        employees_absences (dict): Dictionary of absence lists categorized by key.
        categories (list): List of category keys to iterate over.
    """
    # variables
    h_2 = f"مديرية التربية لولاية أولاد جلال\nمصلحة المستخدمين و التفتيش\nالمؤسسسة : {ecole}-{city}"
    footer_txt = f"\t{city} في :  {first_working_day(month+1, year)}\n\tالمديـــر"
    header_row = {
        "الرقم": 7,
        "الاسم و اللقب": 6,
        "رقم الحساب": 5,
        "الرتبة": 4,
        "من": 3,
        "الي": 2,
        "عدد الايام": 1,
        "ملاحظات": 0
    }
    # create document
    doc: Document = Document()
    configure_document(doc)
   
    try:
            # تحديد الجداول الخاصة بالوثيقة
        absence_lists = [(employees_absences[category], index) for index, category in enumerate(categorys)]

        for absence_list, index_category in absence_lists:
            # Add a new section for each category
            section = doc.add_section() if index_category > 0 else doc.sections[0]            
            section.page_width = Cm(29.7)
            section.page_height = Cm(21)
            section.top_margin = Cm(1)
            section.bottom_margin = Cm(1)
            section.left_margin = Cm(1)
            section.right_margin = Cm(1)
            # Headers and paragraphs
            
            doc.add_paragraph(Logo_header, style="Title").paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.add_paragraph(h_2, style="Paragraph").paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            h_3 = f"كشف الغيابات الشهرية و العطل المرضية لشهر {months[month]}{year} / {categorys[index_category]}"
            doc.add_paragraph(h_3, style="Title").paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Create a new table with the "TableArabic" style
            table = doc.add_table(rows=len(absence_list) + 1, cols=len(header_row), style="TableArabic")
            for key, column_index in header_row.items():    table.cell(0, column_index).text = key
            
            # Fill the table with absence data
            populate_table_with_data(table, absence_list, header_row)
            
            # Define tab stops for the footer
            footer = doc.add_paragraph(footer_txt, style="Paragraph")
            tab_stops = footer.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Cm(24), WD_TAB_ALIGNMENT.CENTER) 

        # Ask user for save path
        file_path, _ = QFileDialog.getSaveFileName(None, "اختر مسار الحفظ", output_doc_name, "ملفات Word (*.docx)")

        # Save the document
        if file_path:
            doc.save(file_path)
            QMessageBox.information(None, "تم التحديث", f"تم تحديث {output_doc_name} بنجاح")

    except Exception as e:
        QMessageBox.warning(None, "خطأ في الحفظ", f"حدث خطأ: {e}")


def populate_table_with_data(table, absence_list, header_row):
    """
    Populate the table with absence data, merging cells for duplicate names.
    
    Args:
        table (Table): The table to populate.
        absence_list (list): List of absence records.
        header_row (dict): Dictionary of header names and their column indices.
    """
    try:    
        row_merged = 0

        for index, data_row in enumerate(absence_list, start=1):
            new_row = table.rows[index].cells
            old_row = table.rows[index - 1]

            if old_row.cells[header_row["الاسم و اللقب"]].text != data_row["name"]:
                # Fill new row with absence data
                fill_data_row(new_row, data_row, header_row, index)
            else:
                # Merge rows for duplicate names
                merge_data_rows(old_row, new_row, data_row, header_row)
                row_merged += 1
    except Exception as e: QMessageBox.warning("populate_table_with_data", "خطأ في تعيين العناصر", str(e))

def fill_data_row(row_cells, data_row, header_row, index):
    """
    Fill a row with data for a single employee.
    
    Args:
        row_cells (list): List of cell objects in the row.
        data_row (dict): Data for the current row.
        header_row (dict): Dictionary of header names and their column indices.
        index (int): Row index.
    """
    try:
        row_cells[header_row["الرقم"]].text = f"{index:2d}"
        row_cells[header_row["الاسم و اللقب"]].text = data_row["name"]
        row_cells[header_row["رقم الحساب"]].text = data_row["ccp"]
        row_cells[header_row["الرتبة"]].text = data_row["grade"]
        row_cells[header_row["من"]].text = data_row["start_date"]
        row_cells[header_row["الي"]].text = data_row["end_date"]
        row_cells[header_row["عدد الايام"]].text = str(data_row["days"])
        row_cells[header_row["ملاحظات"]].text = data_row["motif"] if data_row["motif"] else ""
    except Exception as e: QMessageBox.warning("fill_data_row", "خطأ في تعيين العناصر", str(e))
    
def merge_data_rows(old_row, new_row, data_row, header_row):
    """
    Merge cells in rows with duplicate employee names and update the data.
    
    Args:
        old_row (Row): The previous row to merge with.
        new_row (Row): The current row to merge.
        data_row (dict): Data for the current row.
        header_row (dict): Dictionary of header names and their column indices.
    """
    try:
        old_row.cells[header_row["الرقم"]].merge(new_row[header_row["الرقم"]])
        new_row[header_row["الرقم"]].text = old_row.cells[header_row["الرقم"]].text

        old_row.cells[header_row["الاسم و اللقب"]].merge(new_row[header_row["الاسم و اللقب"]])
        new_row[header_row["الاسم و اللقب"]].text = old_row.cells[header_row["الاسم و اللقب"]].text

        old_row.cells[header_row["رقم الحساب"]].merge(new_row[header_row["رقم الحساب"]])
        new_row[header_row["رقم الحساب"]].text = old_row.cells[header_row["رقم الحساب"]].text

        old_row.cells[header_row["الرتبة"]].merge(new_row[header_row["الرتبة"]])
        new_row[header_row["الرتبة"]].text = old_row.cells[header_row["الرتبة"]].text

        new_row[header_row["من"]].text = data_row["start_date"]
        new_row[header_row["الي"]].text = data_row["end_date"]
        new_row[header_row["عدد الايام"]].text = str(int(old_row.cells[header_row["عدد الايام"]].text) + data_row["days"])
        new_row[header_row["ملاحظات"]].text = data_row["motif"] if data_row["motif"] else ""
    except Exception as e: QMessageBox.warning("merge_data_rows", "خطأ في تعيين العناصر", str(e))
    

def first_working_day(month: int, year: int):


    # Calculate the first day of the next month
    if month == 12:
        next_month = datetime(year + 1, 1, 1)
    else:
        next_month = datetime(year, month + 1, 1)
    
    # Find the first workday (Sunday to Thursday) of the next month
    while next_month.weekday() in [4, 5]:  # 4 is Friday, 5 is Saturday
        next_month += timedelta(days=1)
    
    return next_month.strftime("%d/%m/%Y")


# ====================================================================================================================================


def generate_notices(output_doc_name: str, month: int, year: int, employees_absences: list) -> None:
    
    doc: Document = Document()
    configure_document(doc)
    
    month_name: str = months[month]
    employees_dict = {}

    try:
        # Group absences by employee
        for employee in employees_absences:
            employee_name = employee["name"]
            employee_ccp = employee["ccp"]
            employee_grade = employee["grade"]
            start_date = employee["start_date"]
            end_date = employee["end_date"]
            days_absent = employee["days"]

            if employee_name not in employees_dict:
                employees_dict[employee_name] = {
                    "ccp": employee_ccp,
                    "grade": employee_grade,
                    "absences": []
                }
            employees_dict[employee_name]["absences"].append((start_date, end_date, days_absent))

        notice_count: int = 0  # Counter to keep track of the number of notices in the current section

        # Iterate over employees_dict to create the document
        for employee_name, details in employees_dict.items():
            employee_ccp = details["ccp"]
            employee_grade = details["grade"]
            total_days_absent = sum(absence[2] for absence in details["absences"])

            if  notice_count % 2 == 0:  # Check if the current section already has 2 notices
                section = doc.add_section() if notice_count > 0 else doc.sections[0]
                section.page_width = Cm(21)
                section.page_height = Cm(29.7)
                section.top_margin = Cm(1)
                section.bottom_margin = Cm(1)
                section.left_margin = Cm(1)
                section.right_margin = Cm(1)

            # Logo_header: 
            doc.add_paragraph(Logo_header, style="Title").paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # Header_2:
            header_2_txt = f'مديرية التربية لولاية أولاد جلال\tمن المدير\nالمؤسسة: {ecole}\tإلى السيدة  :  {employee_name}\n{city}\tالرتبة   : {employee_grade}\nالرقم      / {year}\tرقم الحساب : {employee_ccp}'
            header_2 = doc.add_paragraph(header_2_txt, style="Paragraph")
            tab_stops = header_2.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) 
            # title:
            doc.add_paragraph('إشعار بالخصم', style="Title").paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # body:
            paragraph_1 = '    نظر لغيابكم عن العمل  \n '
            absence_details = ""
            for start_date, end_date, days_absent in details["absences"]:
                absence_details += f'من {start_date} إلى {end_date}\n'
            paragraph_2 = absence_details
            paragraph_3 = f'نعلمكم بأنه سيتم خصم   {total_days_absent} يوم )أيام( ، من مرتب شهر {month_name} {year}'

            doc.add_paragraph(paragraph_1 + paragraph_2 + paragraph_3, style="Paragraph").paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            # Footer:
            current_date: str = date.today().strftime("%d-%m-%Y")
            footer_txt = f'المعني\t{city} في : {current_date}\n\tالمدير'
            footer = doc.add_paragraph(footer_txt, style="Paragraph")
            tab_stops = footer.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) 
            # line:
            doc.add_paragraph('\n'+'-'*120).paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            notice_count += 1  # Increment the counter after adding a notice

        file_path, _ = QFileDialog.getSaveFileName(None, "اختر مسار الحفظ", output_doc_name, "ملفات Word (*.docx)")
        if file_path:
            doc.save(file_path)
            QMessageBox.information(None, "تم التحديث", f"تم تحديث {output_doc_name} بنجاح")
    
    except Exception as e:
        QMessageBox.warning(None, "خطأ", f"حدث خطأ: {e}")


